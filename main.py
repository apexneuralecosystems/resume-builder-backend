import os
import json
import tempfile
from pathlib import Path

import httpx
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv

load_dotenv()

app = FastAPI(title="Resume Builder API")

_raw_origins = os.getenv(
    "ALLOWED_ORIGINS",
    "https://resume.builder.apexneural.com,http://localhost:5173,http://localhost:3000,http://localhost",
)
_allowed_origins = [o.strip() for o in _raw_origins.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"


# ── Text extraction ────────────────────────────────────────────────────────────

def extract_text_pdf(file_path: str) -> str:
    import pdfplumber
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def extract_text_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text_doc(file_path: str) -> str:
    import mammoth
    with open(file_path, "rb") as f:
        result = mammoth.extract_raw_text(f)
    return result.value.strip()


def extract_text(file_path: str, content_type: str, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf" or "pdf" in content_type:
            return extract_text_pdf(file_path)
        elif ext == ".docx" or "openxmlformats" in content_type:
            return extract_text_docx(file_path)
        elif ext == ".doc" or "msword" in content_type:
            return extract_text_doc(file_path)
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not extract text from file: {str(e)}")


# ── Prompts ────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are a precise resume parser. Your job is to read the resume text and extract EVERY piece of information faithfully — do NOT invent, guess, or fill in content that is not present in the resume. Return one valid JSON object. No markdown, no explanation, no code fences — only raw JSON.

═══════════════════════════════════════
STRICT SECTION MAPPING — CRITICAL RULE
═══════════════════════════════════════
You must extract each field ONLY from its matching section in the resume:
  • "education"   ← ONLY from Education / Academic / Qualifications section
  • "experience"  ← ONLY from Work Experience / Employment / Career History section
  • "projects"    ← ONLY from Projects / Portfolio / Work section
  • "skills"      ← ONLY from Skills / Technical Skills / Technologies section
  • "certifications" ← ONLY from Certifications / Licenses / Awards section
  • "interests"   ← from Interests / Hobbies / Passions section (if present)
  • "techStack"   ← from Skills / Tech Stack / Tools section
NEVER copy experience bullet points into projects, or education into experience, etc.

═══════════════════════════
COMPLETENESS — CRITICAL RULE
═══════════════════════════
Extract EVERY entry in each section. Do NOT drop entries:
  • If the resume lists 4 degrees → all 4 must appear in "education"
  • If the resume lists 7 jobs → all 7 must appear in "experience"
  • If the resume lists 10 projects → include all 10 in "projects"
  • If a job has 6 achievement bullets → include all 6 in "highlights"

═══════════════
JSON STRUCTURE
═══════════════
{
  "id": "short unique id e.g. usr-001",
  "name": "Full Name from resume header",
  "role": "Most recent or primary job title",
  "email": "email address if present",
  "phone": "phone number if present",
  "location": "City, Country/State if present",
  "website": "personal site URL if present",
  "linkedIn": "LinkedIn URL if present",
  "twitter": "Twitter/X URL if present",
  "github": "GitHub URL if present",

  "bio": "Write a concise 2-3 sentence professional summary in first person based ONLY on what is in the resume. Do not fabricate facts.",
  "aboutMe": "Same as bio but 3-4 sentences, highlighting measurable impact and unique skills found in the resume.",
  "company": "Name of current or most recent employer",
  "yearsExperience": "Total years of experience as a number string e.g. '6'",

  "education": "ALL education entries, one per line, EXACTLY in this format:\n<Institution Name> — <Degree>, <Specialization if any>, <Graduation Year>, <CGPA/Grade if present>\nFor multiple degrees put each on its own line separated by \\n. Example:\nIIT Delhi — B.Tech, Computer Science, 2021, CGPA: 8.7\nXYZ School — HSC, PCM, 2017, 94%",

  "skills": [
    {"name": "Skill name exactly as written in resume", "level": 85}
  ],
  "expertise": ["Top domain expertise areas from resume"],
  "specializations": ["Specific specializations from resume"],
  "interests": ["Professional interests from resume if listed"],

  "techStack": "Group technologies from the resume by category. Use EXACTLY this multi-line format:\nCategory Name\n\n• Technology 1\n• Technology 2\n• Technology 3\n\nNext Category\n\n• Technology A\n• Technology B\nOnly include technologies actually mentioned in the resume.",

  "certifications": [
    {
      "name": "Exact certification name from resume",
      "url": "certification URL if present, else empty string"
    }
  ],

  "projects": [
    {
      "title": "Project title exactly as in resume",
      "description": "Description of what was built, problem solved, and impact — from resume text only",
      "technology": "Technologies listed for this project",
      "link": "Project URL if present"
    }
  ],

  "experience": [
    {
      "role": "Job title exactly as written",
      "company": "Company name exactly as written",
      "period": "Date range as written e.g. Jan 2021 – Mar 2023",
      "location": "Job location if mentioned",
      "type": "Full-time / Part-time / Contract / Internship if mentioned",
      "highlights": [
        "EVERY bullet point / achievement from this role — copy faithfully, include metrics and numbers"
      ]
    }
  ],

  "caseStudies": []
}

═══════════
FIELD RULES
═══════════
1.  education   : One line per degree/qualification using the exact format shown above. Include EVERY degree listed.
2.  experience  : Include EVERY job. For highlights include ALL bullet points from that role (not just 2).
3.  projects    : Include ALL projects listed. Do not truncate.
4.  skills      : Max 7, ranked by prominence in resume. Level: 85-95 primary, 70-84 proficient, 55-69 familiar.
5.  techStack   : 2-4 category blocks, only from resume content.
6.  interests   : Max 8. Only if the resume has an interests/hobbies section. Do not invent.
7.  bio/aboutMe : Derived from the resume summary or from actual experience — never fabricated.
8.  certifications: All listed. Use empty string "" for url if no link is present.
9.  Return ONLY the raw JSON. No markdown, no explanation, no extra keys.
"""

JD_PROMPT_ADDITION = """

IMPORTANT — JOB DESCRIPTION PROVIDED:
Tailor ALL resume content to match this job description while staying truthful to the candidate's actual experience:
1. Set "role" to match the target role in the JD
2. Rewrite "bio" and "aboutMe" to highlight how the candidate is a great fit for this specific role
3. Prioritize and reorder "skills" to put JD-relevant skills first with higher levels
4. Emphasize projects and experiences most relevant to the JD requirements
5. Adjust "interests" to align with the industry/domain in the JD
6. In experience highlights, emphasize achievements that demonstrate JD requirements
7. Reorder "techStack" categories to put JD-relevant tech first

JOB DESCRIPTION:
{jd_text}
"""


# ── OpenRouter call ────────────────────────────────────────────────────────────

async def call_openrouter(api_key: str, system_prompt: str, user_content: str) -> dict:
    payload = {
        "model": "openai/gpt-4o",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        "temperature": 0.2,
        "max_tokens": 4096,
        "response_format": {"type": "json_object"},
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": os.getenv("SITE_URL", "https://resume.builder.apexneural.com"),
        "X-Title": "ResumeForge",
    }

    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(OPENROUTER_URL, json=payload, headers=headers)

    if response.status_code != 200:
        raise HTTPException(
            status_code=502,
            detail=f"OpenRouter error {response.status_code}: {response.text[:300]}",
        )

    data = response.json()
    raw = data["choices"][0]["message"]["content"]
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        raise HTTPException(status_code=502, detail="OpenRouter returned invalid JSON")


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    return {"status": "ok", "openrouter_configured": bool(os.getenv("OPENROUTER_API_KEY"))}


@app.post("/api/parse-resume")
async def parse_resume(
    resume: UploadFile = File(...),
    jd: UploadFile = File(None),
):
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail="OPENROUTER_API_KEY is not configured. Please add it to backend/.env",
        )

    # Extract resume text
    suffix = Path(resume.filename or "file.pdf").suffix or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await resume.read())
        tmp_path = tmp.name
    try:
        resume_text = extract_text(tmp_path, resume.content_type or "", resume.filename or "file")
    finally:
        os.unlink(tmp_path)

    if not resume_text.strip():
        raise HTTPException(status_code=422, detail="Could not extract any text from the resume file")

    # Extract JD text if provided
    jd_text = None
    if jd and jd.filename and jd.size and jd.size > 0:
        jd_suffix = Path(jd.filename).suffix or ".pdf"
        with tempfile.NamedTemporaryFile(delete=False, suffix=jd_suffix) as tmp:
            tmp.write(await jd.read())
            jd_tmp_path = tmp.name
        try:
            jd_text = extract_text(jd_tmp_path, jd.content_type or "", jd.filename)
        finally:
            os.unlink(jd_tmp_path)

    # Build prompt
    system_prompt = SYSTEM_PROMPT
    if jd_text and jd_text.strip():
        system_prompt += JD_PROMPT_ADDITION.format(jd_text=jd_text[:3000])

    result = await call_openrouter(
        api_key=api_key,
        system_prompt=system_prompt,
        user_content=f"Parse this resume:\n\n{resume_text[:6000]}",
    )
    return result


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
